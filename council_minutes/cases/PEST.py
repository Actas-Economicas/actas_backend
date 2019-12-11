from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, BooleanField
from ..models import Request
from .case_utils import add_analysis_paragraph


class PEST(Request):

    full_name = 'Práctica estudiantil'

    SUB_P1 = 'P1'
    SUB_P2 = 'P2'
    SUB_P3 = 'P3'
    SUBJECT_CHOICES = (
        (SUB_P1, 'Práctica Estudiantil I'),
        (SUB_P2, 'Práctica Estudiantil II'),
        (SUB_P3, 'Práctica Estudiantil III')
    )
    SUBJECT_INFO = {
        SUB_P1: ('2016762', 3),
        SUB_P2: ('2016763', 6),
        SUB_P3: ('2016764', 9)
    }

    institution = StringField(required=True, display='Institución/Empresa', default='')
    proffesor = StringField(required=True, display='Profesor', default='')
    ins_person = StringField(required=True, display='Encargado Institucion', default='')
    subject = StringField(required=True, choices=SUBJECT_CHOICES,
                          default=SUB_P1, display='Asignatura')
    advance = FloatField(required=True, min_value=0, display='Avance SIA', default=0.0)
    another_practice = BooleanField(
        required=True, display='¿Primera practica?', default=False)
    hours = IntField(required=True, min_value=0, display='Horas Semana', default=0)
    duration = StringField(required=True, display='Duración', default='')
    documentation = BooleanField(
        required=True, display='¿Documentación Completa?', default=True)

    regulation_list = ['008|2008|CSU', '102|2013|CSU', '016|2011|CAC']

    str_cm = [
        'inscribir la asignatura {} ({}) con carga de {} créditos, ',
        'en el periodo {}, a desarrollar en la empresa {}, a cargo del docente ' +
        '{} por parte de la Universidad Nacional de Colombia y {} por parte de la entidad ' +
        '(Artículo 15 {}).',
        'debido a que {} ({}).'
    ]

    str_analysis = [
        'El estudiante {}cumple con el requisito de haber aprobado el ' +
        '70% de los créditos del plan de estudios. SIA: {:0.1f}% de avance en ' +
        'los créditos exigidos del plan de estudios.',
        'El estudiante {}ha cursado otra de las asignaturas con ' +
        'el nombre Práctica Estudiantil.',
        'Requisitos: Pertinencia, objetivos, alcance, empresa {}, duración: {} ' +
        'horas/semana durante {}, costos, descripción de actividades ' +
        'a cargo de un profesor de la Facultad: {}, porcentajes de evaluación definidos ' +
        '(Artículo 3 del {}).',
        'Documentación {}cumple con requisitos: Formato está completamente diligenciado, ' +
        'adjunta copia del Acuerdo firmado, ' +
        'adjunta el recibido de la carta de presentación de la Universidad, ' +
        'están fijados los porcentajes de evaluación.'
    ]

    str_pcm = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        # pylint: disable=no-member
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            self.get_approval_status_display().upper() + ' ').font.bold = True
        self.add_text(
            paragraph, self.is_affirmative_response_approval_status())

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.add_analysis())
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        self.add_text(
            paragraph, self.is_affirmative_response_advisor_response())

    def add_text(self, paragraph, affirmative):
        code, _credits = self.SUBJECT_INFO[self.subject]
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0].format(
            self.get_subject_display(), code, _credits))

        if affirmative:
            paragraph.add_run(self.str_cm[1].format(
                self.academic_period, self.institution,
                self.proffesor, self.ins_person,
                self.regulations[self.regulation_list[0]][0]
            ))
        else:
            paragraph.add_run(self.str_cm[2].format(
                self.council_decision,
                self.regulations[self.regulation_list[1]][0]
            ))

    def add_analysis(self):
        analysis = []
        modifier = '' if self.advance >= 70 else 'no '
        analysis.append(self.str_analysis[0].format(modifier, self.advance))

        modifier = '' if self.another_practice else 'no '
        analysis.append(self.str_analysis[1].format(modifier))

        analysis.append(self.str_analysis[2].format(
            self.institution, self.hours, self.duration, self.proffesor,
            self.regulations[self.regulation_list[2]][0]))

        modifier = '' if self.documentation else 'no '
        self.str_analysis[3].format(self.str_analysis[3].format(modifier))

        return analysis + self.extra_analysis
