from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, FloatField
from ..models import Request, Professor
from .case_utils import add_analysis_paragraph


class ADIC(Request):

    full_name = 'Adición de codirector'

    node = StringField(required=True, display='Perfil')
    title = StringField(required=True, display='Título de Tesis/Trabajo Final')
    council_number = StringField(
        required=True, max_length=2, default='00', display='# Acta de cancelación')
    council_year = StringField(
        required=True, min_length=4, max_length=4, display='Año del Acta')
    proffesors = EmbeddedDocumentListField(
        Professor, required=True, display='Docentes')

    regulation_list = []

    str_cm = []
    str_pcm = []

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(self.str_council_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True

    def pcm(self, docx):
        add_analysis_paragraph(docx, self.extra_analysis)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(self.str_answer + ' ').font.bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
