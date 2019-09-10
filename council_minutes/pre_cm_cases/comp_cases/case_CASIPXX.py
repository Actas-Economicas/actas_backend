from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CASIPXX():

    count = 0

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS(request, docx, redirected=False):
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis(request, docx)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Answers(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis(request, docx):
        para = docx.add_paragraph()
        para.add_run('Analisis:')
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_1(request, docx)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_2(request, docx)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_3(request, docx)
        CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_extra(request, docx)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_1(request, docx):
        str_in = '1. SIA: Porcentaje de avance en el plan: {}. Número de'
        str_in += 'matrículas: {}. PAPA: {}.'
        docx.add_paragraph(str_in.format(request['pre_cm']['advance'],
                           request['pre_cm']['enrolled_academic_periods'],
                           request['pre_cm']['papa']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_2(request, docx):
        str_in = '2. SIA: Créditos disponibles: {}.'
        docx.add_paragraph(str_in.format(request['pre_cm']['available']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_S(docx, subject):
        str_in = '{}. SIA: Al aprobar la cancelación de la asignatura {} ({}) '
        str_in += ' el estudiante quedaría con {} créditos inscritos.'
        docx.add_paragraph(str_in.format(subject['number'], subject['code'],
                           subject['subject'], subject['remaining']))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_3(request, docx):
        CASIPXX.count = 2
        for subject in request['detail_cm']['subjects']:
            CASIPXX.count = CASIPXX.count + 1
            subject['number'] = str(CASIPXX.count)
            current_credits = int(request['pre_cm']['current_credits'])
            subject_credits = int(subject['credits'])
            subject['remaining'] = current_credits - subject_credits
            CASIPXX.case_CANCELACION_DE_ASIGNATURAS_Analysis_S(docx, subject)

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Analysis_extra(request, docx):
        for analysis in request['pre_cm']['extra_analysis']:
            CASIPXX.count = CASIPXX.count + 1
            str_in = '{}. {}.'
            docx.add_paragraph(str_in.format(CASIPXX.count, analysis))

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers(request, docx):
        if request['approval_status'] == 'RC':
            pass
        else:
            pass

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers_RC(request, docx):
        pass

    @staticmethod
    def case_CANCELACION_DE_ASIGNATURAS_Answers_NRC(request, docx):
        pass
